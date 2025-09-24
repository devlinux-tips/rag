"""
Document text extraction system supporting multiple file formats.
Provides reliable text extraction from PDF, DOCX, and plain text files
with proper encoding handling and error recovery.
"""

from dataclasses import dataclass
from pathlib import Path
from typing import Any, Protocol, runtime_checkable

from bs4 import BeautifulSoup
from docx import Document
from pypdf import PdfReader

from ..utils.logging_factory import (
    get_system_logger,
    log_component_end,
    log_component_start,
    log_data_transformation,
    log_decision_point,
    log_error_context,
    log_performance_metric,
)


@dataclass
class ExtractionResult:
    """Result of document text extraction."""

    text: str
    character_count: int
    extraction_method: str
    encoding_used: str | None = None
    pages_processed: int | None = None
    error_details: str | None = None


@dataclass
class ExtractionConfig:
    """Configuration for document extraction."""

    supported_formats: list[str]
    text_encodings: list[str]
    max_file_size_mb: int = 50
    enable_logging: bool = True


@runtime_checkable
class FileSystemProvider(Protocol):
    """Protocol for file system operations to enable testing."""

    def file_exists(self, file_path: Path) -> bool:
        """Check if file exists."""
        ...

    def get_file_size_mb(self, file_path: Path) -> float:
        """Get file size in MB."""
        ...

    def open_binary(self, file_path: Path) -> bytes:
        """Open file in binary mode."""
        ...

    def open_text(self, file_path: Path, encoding: str) -> str:
        """Open file in text mode with specified encoding."""
        ...


@runtime_checkable
class ConfigProvider(Protocol):
    """Protocol for configuration access."""

    def get_extraction_config(self) -> dict[str, Any]:
        """Get extraction configuration."""
        ...


@runtime_checkable
class LoggerProvider(Protocol):
    """Protocol for logging operations."""

    def info(self, message: str) -> None:
        """Log info message."""
        ...

    def error(self, message: str) -> None:
        """Log error message."""
        ...


# ================================
# PURE BUSINESS LOGIC FUNCTIONS
# ================================


def validate_file_format(file_path: Path, supported_formats: list[str]) -> tuple[bool, str | None]:
    """
    Validate if file format is supported.

    Args:
        file_path: Path to document file
        supported_formats: List of supported file extensions

    Returns:
        Tuple of (is_valid, error_message)
    """
    logger = get_system_logger()
    logger.trace("file_extractor", "validate_file_format", f"Validating format for {file_path.name}")

    suffix = file_path.suffix.lower()
    logger.debug("file_extractor", "validate_file_format", f"File extension: {suffix}")

    if suffix not in supported_formats:
        error_msg = f"Unsupported format: {suffix}. Supported: {supported_formats}"
        logger.warning("file_extractor", "validate_file_format", error_msg)
        return False, error_msg

    logger.debug("file_extractor", "validate_file_format", f"Format validation passed for {suffix}")
    return True, None


def extract_text_from_pdf_binary(pdf_binary: bytes) -> tuple[str, int]:
    """
    Extract text from PDF binary data.
    Pure function with no side effects.

    Args:
        pdf_binary: PDF file content as bytes

    Returns:
        Tuple of (extracted_text, pages_processed)

    Raises:
        Exception: If PDF processing fails
        ImportError: If pypdf is not available
    """
    logger = get_system_logger()
    log_component_start("pdf_extractor", "extract_text_from_pdf_binary", binary_size=len(pdf_binary))

    from io import BytesIO

    text_content = []

    try:
        with BytesIO(pdf_binary) as pdf_stream:
            pdf_reader = PdfReader(pdf_stream)
            pages_processed = len(pdf_reader.pages)
            logger.debug("pdf_extractor", "extract_text_from_pdf_binary", f"PDF has {pages_processed} pages")

            for page_num in range(pages_processed):
                logger.trace("pdf_extractor", "extract_text_from_pdf_binary", f"Processing page {page_num + 1}")
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text.strip():
                    text_content.append(text.strip())
                    logger.trace(
                        "pdf_extractor", "extract_text_from_pdf_binary", f"Page {page_num + 1}: {len(text)} chars"
                    )

        full_text = "\n\n".join(text_content)
        log_data_transformation("pdf_extractor", "join_pages", f"{pages_processed} pages", f"text[{len(full_text)}]")
        log_component_end(
            "pdf_extractor",
            "extract_text_from_pdf_binary",
            f"Extracted {len(full_text)} chars",
            pages_processed=pages_processed,
            total_chars=len(full_text),
        )
        return full_text, pages_processed

    except Exception as e:
        log_error_context("pdf_extractor", "extract_text_from_pdf_binary", e, {"binary_size": len(pdf_binary)})
        raise


def extract_text_from_docx_binary(docx_binary: bytes) -> str:
    """
    Extract text from DOCX binary data.
    Pure function with no side effects.

    Args:
        docx_binary: DOCX file content as bytes

    Returns:
        Extracted text content

    Raises:
        Exception: If DOCX processing fails
        ImportError: If python-docx is not available
    """
    logger = get_system_logger()
    log_component_start("docx_extractor", "extract_text_from_docx_binary", binary_size=len(docx_binary))

    from io import BytesIO

    text_content = []

    try:
        with BytesIO(docx_binary) as docx_stream:
            doc = Document(docx_stream)
            paragraphs_total = len(doc.paragraphs)
            logger.debug("docx_extractor", "extract_text_from_docx_binary", f"DOCX has {paragraphs_total} paragraphs")

            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
                    logger.trace(
                        "docx_extractor",
                        "extract_text_from_docx_binary",
                        f"Paragraph {i + 1}: {len(paragraph.text)} chars",
                    )

        full_text = "\n\n".join(text_content)
        log_data_transformation(
            "docx_extractor", "join_paragraphs", f"{paragraphs_total} paragraphs", f"text[{len(full_text)}]"
        )
        log_component_end(
            "docx_extractor",
            "extract_text_from_docx_binary",
            f"Extracted {len(full_text)} chars",
            paragraphs_processed=paragraphs_total,
            total_chars=len(full_text),
        )
        return full_text

    except Exception as e:
        log_error_context("docx_extractor", "extract_text_from_docx_binary", e, {"binary_size": len(docx_binary)})
        raise


def extract_text_with_encoding_fallback(text_binary: bytes, encodings: list[str]) -> tuple[str, str]:
    """
    Extract text with multiple encoding fallback.
    Pure function with no side effects.

    Args:
        text_binary: Text file content as bytes
        encodings: List of encodings to try in order

    Returns:
        Tuple of (extracted_text, successful_encoding)

    Raises:
        ValueError: If no encoding succeeds
    """
    logger = get_system_logger()
    log_component_start(
        "text_extractor",
        "extract_text_with_encoding_fallback",
        binary_size=len(text_binary),
        encodings_count=len(encodings),
    )

    for i, encoding in enumerate(encodings):
        logger.trace(
            "text_extractor",
            "extract_text_with_encoding_fallback",
            f"Trying encoding {i + 1}/{len(encodings)}: {encoding}",
        )
        try:
            text = text_binary.decode(encoding).strip()
            log_decision_point(
                "text_extractor",
                "extract_text_with_encoding_fallback",
                f"encoding={encoding}",
                f"success with {len(text)} chars",
            )
            log_component_end(
                "text_extractor",
                "extract_text_with_encoding_fallback",
                f"Decoded with {encoding}",
                successful_encoding=encoding,
                text_length=len(text),
            )
            return text, encoding
        except UnicodeDecodeError as e:
            logger.debug(
                "text_extractor", "extract_text_with_encoding_fallback", f"Encoding {encoding} failed: {str(e)}"
            )
            continue

    error_msg = f"Could not decode text with any supported encoding: {encodings}"
    logger.error("text_extractor", "extract_text_with_encoding_fallback", error_msg)
    raise ValueError(error_msg)


def post_process_extracted_text(text: str) -> str:
    """
    Post-process extracted text for consistency.
    Pure function with no side effects.

    Args:
        text: Raw extracted text

    Returns:
        Processed text content
    """
    if not text:
        return ""

    # Basic text cleaning
    processed_text = text.strip()

    # Normalize whitespace while preserving paragraphs
    lines = processed_text.split("\n")
    cleaned_lines = [line.strip() for line in lines if line.strip()]

    return "\n\n".join(cleaned_lines)


# ================================
# DEPENDENCY INJECTION ORCHESTRATION
# ================================


class DocumentExtractor:
    """Document extractor with dependency injection for testability."""

    def __init__(
        self,
        config_provider: ConfigProvider,
        file_system_provider: FileSystemProvider,
        logger_provider: LoggerProvider | None = None,
    ):
        """Initialize extractor with injected dependencies."""
        logger = get_system_logger()
        log_component_start("document_extractor", "init", has_logger=logger_provider is not None)

        self._config_provider = config_provider
        self._file_system = file_system_provider
        self._logger = logger_provider

        logger.debug("document_extractor", "init", "Loading extraction configuration")
        config_data = config_provider.get_extraction_config()

        if "max_file_size_mb" not in config_data:
            logger.error("document_extractor", "init", "Missing 'max_file_size_mb' in extraction configuration")
            raise ValueError("Missing 'max_file_size_mb' in extraction configuration")
        if "enable_logging" not in config_data:
            logger.error("document_extractor", "init", "Missing 'enable_logging' in extraction configuration")
            raise ValueError("Missing 'enable_logging' in extraction configuration")

        self._config = ExtractionConfig(
            supported_formats=config_data["supported_formats"],
            text_encodings=config_data["text_encodings"],
            max_file_size_mb=config_data["max_file_size_mb"],
            enable_logging=config_data["enable_logging"],
        )

        log_component_end(
            "document_extractor",
            "init",
            "Extraction system initialized",
            supported_formats=self._config.supported_formats,
            max_file_size_mb=self._config.max_file_size_mb,
            encodings_count=len(self._config.text_encodings),
        )

    def extract_text(self, file_path: Path) -> ExtractionResult:
        """
        Extract text from document using dependency injection.

        Args:
            file_path: Path to document file

        Returns:
            ExtractionResult with text and metadata

        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file format is not supported or file too large
        """
        logger = get_system_logger()
        log_component_start("document_extractor", "extract_text", file_path=str(file_path))

        if not self._file_system.file_exists(file_path):
            logger.error("document_extractor", "extract_text", f"File not found: {file_path}")
            raise FileNotFoundError(f"File not found: {file_path}")

        file_size_mb = self._file_system.get_file_size_mb(file_path)
        log_performance_metric("document_extractor", "extract_text", "file_size_mb", file_size_mb)

        if file_size_mb > self._config.max_file_size_mb:
            logger.error(
                "document_extractor",
                "extract_text",
                f"File too large: {file_size_mb:.1f}MB > {self._config.max_file_size_mb}MB",
            )
            raise ValueError(f"File too large: {file_size_mb:.1f}MB > {self._config.max_file_size_mb}MB")

        is_valid, error_msg = validate_file_format(file_path, self._config.supported_formats)
        if not is_valid:
            error_message = error_msg or "File format validation failed"
            logger.error("document_extractor", "extract_text", error_message)
            raise ValueError(error_message)

        self._log_info(f"Extracting text from {file_path}")

        suffix = file_path.suffix.lower()
        log_decision_point("document_extractor", "extract_text", f"file_extension={suffix}", "route_to_extractor")

        try:
            if suffix == ".pdf":
                result = self._extract_pdf(file_path)
            elif suffix == ".docx":
                result = self._extract_docx(file_path)
            elif suffix == ".txt":
                result = self._extract_txt(file_path)
            elif suffix == ".html":
                result = self._extract_html(file_path)
            elif suffix == ".md":
                result = self._extract_txt(file_path)  # Markdown can be processed as text
            else:
                raise ValueError(f"Unsupported format: {suffix}")

            log_component_end(
                "document_extractor",
                "extract_text",
                f"Extracted {result.character_count} chars",
                extraction_method=result.extraction_method,
                character_count=result.character_count,
                file_size_mb=file_size_mb,
            )
            return result

        except Exception as e:
            log_error_context(
                "document_extractor",
                "extract_text",
                e,
                {"file_path": str(file_path), "file_size_mb": file_size_mb, "file_extension": suffix},
            )
            raise

    def _extract_pdf(self, file_path: Path) -> ExtractionResult:
        """Extract text from PDF using pure function."""
        pdf_binary = self._file_system.open_binary(file_path)
        text, pages_processed = extract_text_from_pdf_binary(pdf_binary)
        processed_text = post_process_extracted_text(text)

        self._log_info(f"Extracted {len(processed_text)} characters from {pages_processed} PDF pages")

        return ExtractionResult(
            text=processed_text,
            character_count=len(processed_text),
            extraction_method="PDF",
            pages_processed=pages_processed,
        )

    def _extract_docx(self, file_path: Path) -> ExtractionResult:
        """Extract text from DOCX using pure function."""
        docx_binary = self._file_system.open_binary(file_path)
        text = extract_text_from_docx_binary(docx_binary)
        processed_text = post_process_extracted_text(text)

        self._log_info(f"Extracted {len(processed_text)} characters from DOCX")

        return ExtractionResult(text=processed_text, character_count=len(processed_text), extraction_method="DOCX")

    def _extract_txt(self, file_path: Path) -> ExtractionResult:
        """Extract text from TXT using pure function with encoding fallback."""
        text_binary = self._file_system.open_binary(file_path)
        text, encoding_used = extract_text_with_encoding_fallback(text_binary, self._config.text_encodings)
        processed_text = post_process_extracted_text(text)

        self._log_info(f"Successfully read TXT file with {encoding_used} encoding")
        self._log_info(f"Extracted {len(processed_text)} characters from TXT")

        return ExtractionResult(
            text=processed_text,
            character_count=len(processed_text),
            extraction_method="TXT",
            encoding_used=encoding_used,
        )

    def _extract_html(self, file_path: Path) -> ExtractionResult:
        """Extract text from HTML using BeautifulSoup with encoding fallback."""
        html_binary = self._file_system.open_binary(file_path)
        html_content, encoding_used = extract_text_with_encoding_fallback(html_binary, self._config.text_encodings)

        # Parse HTML and extract text content
        soup = BeautifulSoup(html_content, "html.parser")

        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()

        # Extract text and clean up
        text = soup.get_text()
        processed_text = post_process_extracted_text(text)

        self._log_info(f"Successfully parsed HTML file with {encoding_used} encoding")
        self._log_info(f"Extracted {len(processed_text)} characters from HTML")

        return ExtractionResult(
            text=processed_text,
            character_count=len(processed_text),
            extraction_method="HTML",
            encoding_used=encoding_used,
        )

    def _log_info(self, message: str) -> None:
        """Log info message if logger available."""
        if self._logger and self._config.enable_logging:
            self._logger.info(message)

    def _log_error(self, message: str) -> None:
        """Log error message if logger available."""
        if self._logger and self._config.enable_logging:
            self._logger.error(message)


# ================================
# CONVENIENCE FUNCTIONS
# ================================


def extract_document_text(
    file_path: str | Path,
    config_provider: ConfigProvider | None = None,
    file_system_provider: FileSystemProvider | None = None,
) -> str:
    """
    Convenience function for document text extraction.

    Args:
        file_path: Path to document file
        config_provider: Optional config provider (uses production if None)
        file_system_provider: Optional file system provider (uses production if None)

    Returns:
        Extracted text content
    """
    from .extractors_providers import create_config_provider, create_file_system_provider

    if isinstance(file_path, str):
        file_path = Path(file_path)

    # Use injected providers or create defaults
    config = config_provider or create_config_provider()
    file_system = file_system_provider or create_file_system_provider()

    extractor = DocumentExtractor(config, file_system)
    result = extractor.extract_text(file_path)

    return result.text
